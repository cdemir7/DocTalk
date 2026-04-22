from abc import ABC, abstractmethod
from pathlib import Path
import re

import pymupdf4llm
from docx import Document as DocxDocument

from backend.config import settings


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, filePath: Path) -> str:
        ...


class PdfParser(DocumentParser):
    def parse(self, filePath: Path) -> str:
        return pymupdf4llm.to_markdown(str(filePath))


class WordParser(DocumentParser):
    def parse(self, filePath: Path) -> str:
        doc = DocxDocument(str(filePath))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)


class TxtParser(DocumentParser):
    def parse(self, filePath: Path) -> str:
        raw = filePath.read_text(encoding="utf-8")
        normalized = raw.replace("\r\n", "\n").replace("\r", "\n")
        return re.sub(r"\n{3,}", "\n\n", normalized).strip()


_parsers: dict[str, DocumentParser] = {
    ".pdf": PdfParser(),
    ".docx": WordParser(),
    ".txt": TxtParser(),
}


def parseDocument(filePath: Path) -> Path:
    suffix = filePath.suffix.lower()

    parser = _parsers.get(suffix)
    if parser is None:
        raise ValueError(f"Desteklenmeyen dosya tipi: {suffix}")

    parsedText = parser.parse(filePath)
    outputPath = settings.parsedResourcesDirectory / f"{filePath.stem}_parsed.md"
    outputPath.write_text(parsedText, encoding="utf-8")

    return outputPath
